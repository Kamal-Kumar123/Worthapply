"""Resume parser — extracts text from PDF and DOCX files."""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import BinaryIO

logger = logging.getLogger(__name__)


def parse_pdf(source: str | Path | BinaryIO) -> str:
    """Extract text from a PDF file or file-like object."""
    from PyPDF2 import PdfReader

    try:
        if isinstance(source, (str, Path)):
            reader = PdfReader(str(source))
        else:
            reader = PdfReader(source)

        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as exc:
        logger.error("PDF parse error: %s", exc)
        return ""


def parse_docx(source: str | Path | BinaryIO) -> str:
    """Extract text from a DOCX file or file-like object."""
    from docx import Document

    try:
        if isinstance(source, (str, Path)):
            doc = Document(str(source))
        else:
            doc = Document(source)

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return "\n\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX parse error: %s", exc)
        return ""


def parse_resume(source: str | Path | BinaryIO, filename: str = "") -> str:
    """Parse a resume file, auto-detecting format from extension or filename.

    Args:
        source: File path or file-like object.
        filename: Original filename (used for format detection when source is a stream).

    Returns:
        Extracted text, or empty string on failure.
    """
    if isinstance(source, (str, Path)):
        ext = Path(source).suffix.lower()
    else:
        ext = Path(filename).suffix.lower() if filename else ""

    if ext == ".pdf":
        return parse_pdf(source)
    elif ext in (".docx", ".doc"):
        return parse_docx(source)
    elif ext in (".txt", ".md"):
        if isinstance(source, (str, Path)):
            return Path(source).read_text(encoding="utf-8")
        else:
            raw = source.read()
            return raw.decode("utf-8") if isinstance(raw, bytes) else raw
    else:
        logger.warning("Unsupported resume format: %s", ext or "(unknown)")
        return ""
